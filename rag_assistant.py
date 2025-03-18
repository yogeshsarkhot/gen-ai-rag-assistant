import streamlit as st
import os
import pandas as pd
import shutil
import time
import re

from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_ollama import OllamaLLM
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from PyPDF2 import PdfReader
from docx import Document
from langchain_core.documents import Document as LangChainDocument

# Initialize session state
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "qa_history" not in st.session_state:
    st.session_state.qa_history = []
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Persistent directory for Chroma
CHROMA_PERSIST_DIR = "./chroma_db"
COLLECTION_NAME = "rag_collection"

# Embedding model
embeddings = OllamaEmbeddings(model="nomic-embed-text")

# Initialize or load vector store
def initialize_vector_store(force_reset=False):
    if force_reset and os.path.exists(CHROMA_PERSIST_DIR):
        shutil.rmtree(CHROMA_PERSIST_DIR, ignore_errors=True)
    
    if os.path.exists(CHROMA_PERSIST_DIR):
        try:
            vector_store = Chroma(
                persist_directory=CHROMA_PERSIST_DIR,
                embedding_function=embeddings,
                collection_name=COLLECTION_NAME
            )
            vector_store._collection.count()
            return vector_store
        except Exception as e:
            st.warning(f"Failed to load vector store: {str(e)}. Resetting directory.")
            shutil.rmtree(CHROMA_PERSIST_DIR, ignore_errors=True)
    
    return None

# Streamlit UI
st.title("Local free, secure and personalized Generative AI RAG assistant")

# File uploader
uploaded_files = st.file_uploader("Upload files", type=["txt", "pdf", "docx", "xlsx"], accept_multiple_files=True)

def preprocess_text(text):
    """Normalize text for better embeddings and matching."""
    text = re.sub(r'\s+', ' ', text)  # Collapse multiple spaces/newlines into single space
    text = text.strip()  # Remove leading/trailing whitespace
    return text.lower()  # Convert to lowercase for consistent matching

def process_file(file, file_type):
    """Process a file and return LangChain documents with correct metadata."""
    temp_file = f"temp_file_{file.name}"
    with open(temp_file, "wb") as f:
        f.write(file.getbuffer())
    
    original_filename = file.name
    
    if file_type == "txt":
        loader = TextLoader(temp_file)
        documents = loader.load()
        for doc in documents:
            doc.page_content = preprocess_text(doc.page_content)
            doc.metadata = {"source": original_filename}
    elif file_type == "pdf":
        reader = PdfReader(temp_file)
        text = "".join(page.extract_text() + " " for page in reader.pages if page.extract_text())
        text = preprocess_text(text)
        documents = [LangChainDocument(page_content=text, metadata={"source": original_filename})]
    elif file_type == "docx":
        doc = Document(temp_file)
        text = " ".join(para.text for para in doc.paragraphs if para.text.strip())  # Join paragraphs with space
        text = preprocess_text(text)
        documents = [LangChainDocument(page_content=text, metadata={"source": original_filename})]
    elif file_type == "xlsx":
        df = pd.read_excel(temp_file)
        text = " ".join(df.astype(str).agg(' '.join, axis=1))  # Convert to string and join rows with space
        text = preprocess_text(text)
        documents = [LangChainDocument(page_content=text, metadata={"source": original_filename})]
    
    os.remove(temp_file)
    return documents

# Process uploaded files
if uploaded_files:
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    
    for uploaded_file in uploaded_files:
        if uploaded_file.name not in st.session_state.processed_files:
            try:
                file_type = uploaded_file.name.split(".")[-1].lower()
                documents = process_file(uploaded_file, file_type)
                chunks = text_splitter.split_documents(documents)
                for chunk in chunks:
                    chunk.metadata = {"source": uploaded_file.name}

                if st.session_state.vector_store is None:
                    st.session_state.vector_store = Chroma.from_documents(
                        documents=chunks,
                        embedding=embeddings,
                        collection_name=COLLECTION_NAME,
                        persist_directory=CHROMA_PERSIST_DIR
                    )
                    st.success(f"Created new vector store with '{uploaded_file.name}'")
                else:
                    st.session_state.vector_store.add_documents(chunks)
                    st.success(f"Added '{uploaded_file.name}' to vector store")
                
                st.session_state.processed_files.add(uploaded_file.name)
            except Exception as e:
                st.error(f"Error processing '{uploaded_file.name}': {str(e)}")

# Display processed files
if st.session_state.processed_files:
    st.subheader("Processed Files")
    for file_name in st.session_state.processed_files:
        st.write(f"- {file_name}")

# Question input
st.subheader("Ask a Question")
with st.form(key="question_form", clear_on_submit=True):
    question = st.text_input("Ask a question about the documents:")
    submit_button = st.form_submit_button(label="Submit")

if submit_button and question:
    if st.session_state.vector_store is None:
        st.session_state.vector_store = initialize_vector_store()
    
    if st.session_state.vector_store:
        llm = OllamaLLM(model="llama3")
        
        retriever = st.session_state.vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 5}
        )

        # Debug retrieved documents
        with st.expander("Debug Information"):
            count = st.session_state.vector_store._collection.count()
            st.write(f"Total chunks in vector store: {count}")
            retrieved_docs = retriever.invoke(question)
            st.write("Retrieved Documents:")
            for i, doc in enumerate(retrieved_docs):
                source = doc.metadata.get("source", "Unknown")
                st.write(f"Doc {i+1} from '{source}': {doc.page_content[:200]}...")

        # Prompt for stuffing all documents
        prompt = PromptTemplate(
            input_variables=["context", "question"],
            template=(
                "You are a precise assistant answering based solely on the provided document content. "
                "Extract the answer directly from the context below. If the context doesn’t contain the answer, respond only with 'Unable to answer based on documents.'\n\n"
                "Context: {context}\n\n"
                "Question: {question}\n\n"
                "Answer:"
            )
        )

        # QA chain with stuff
        try:
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=retriever,
                return_source_documents=True,
                chain_type_kwargs={
                    "prompt": prompt,
                    "document_variable_name": "context"
                }
            )

            with st.spinner("Generating answer..."):
                result = qa_chain.invoke({"query": question})
                answer = result["result"].strip()
                source_docs = result.get("source_documents", [])

                # Determine contributing sources by matching the exact answer or key phrase
                contributing_sources = set()
                normalized_answer = answer.lower().strip()
                unable_to_answer = "unable to answer based on documents" in normalized_answer
                core_answer = None  # Initialize core_answer outside the if block

                if not unable_to_answer:
                    # Extract the core answer by removing preamble and quotes
                    core_answer_match = re.search(r'["“](.*?)["”]$', normalized_answer)
                    core_answer = core_answer_match.group(1) if core_answer_match else normalized_answer
                    core_answer = re.sub(r'^according to.*?:\s*', '', core_answer).strip()

                    # Debug the extracted core answer and document contents
                    with st.expander("Source Matching Debug"):
                        st.write(f"Core Answer to Match: '{core_answer}'")
                        for i, doc in enumerate(source_docs):
                            st.write(f"Doc {i+1} Content (first 200 chars): '{doc.page_content[:200]}...'")
                            st.write(f"Doc {i+1} Source: '{doc.metadata.get('source', 'Unknown')}'")

                    # Look for the core answer in each document
                    for doc in source_docs:
                        doc_content = doc.page_content.lower()  # Already normalized in preprocess_text
                        if core_answer in doc_content:
                            contributing_sources.add(doc.metadata.get("source", "Unknown"))

                    # If no exact match, use significant keywords as fallback
                    if not contributing_sources:
                        answer_keywords = set(re.split(r'\W+', core_answer)) - set(["the", "is", "of", "in", "a", "an", "according", "to"])
                        significant_keywords = {kw for kw in answer_keywords if len(kw) > 3}
                        for doc in source_docs:
                            doc_content = doc.page_content.lower()
                            if any(keyword in doc_content for keyword in significant_keywords):
                                contributing_sources.add(doc.metadata.get("source", "Unknown"))

                # Debug LLM input and output
                with st.expander("LLM Debug"):
                    context = "\n".join(doc.page_content for doc in source_docs)
                    st.write(f"Context passed to LLM:\n{context[:500]}..." if len(context) > 500 else context)
                    st.write(f"Final Answer: '{answer}'")
                    st.write(f"Core Answer Extracted: '{core_answer if core_answer else 'N/A'}'")
                    st.write(f"Contributing Sources: {contributing_sources}")

                history_entry = {
                    "question": question,
                    "answer": answer,
                    "sources": list(contributing_sources) if not unable_to_answer else []
                }

                st.session_state.qa_history.append(history_entry)
        except Exception as e:
            st.error(f"Error initializing QA chain or generating answer: {str(e)}")
    else:
        st.error("No documents loaded. Please upload files first.")

# Display history
st.subheader("Question and Answer History")
for i, qa in enumerate(st.session_state.qa_history):
    st.write(f"**Q{i+1}:** {qa['question']}")
    st.write(f"**A{i+1}:** {qa['answer']}")
    if qa["sources"]:
        st.write(f"**Sources:** {', '.join(qa['sources'])}")
    else:
        st.write("**Sources:** None")
    st.write("---")

# Reset button
if st.button("Reset Knowledge Base"):
    try:
        st.session_state.vector_store = initialize_vector_store(force_reset=True)
        st.session_state.processed_files = set()
        st.session_state.qa_history = []
        time.sleep(1)
        st.success("Knowledge base reset successfully!")
        st.rerun()
    except Exception as e:
        st.error(f"Error resetting: {str(e)}")
